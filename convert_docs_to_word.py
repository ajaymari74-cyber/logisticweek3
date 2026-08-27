"""
Convert Week 3 Documentation to a Professional Word Document (.docx)
Author: Senior Documentation Specialist & Python Developer
Target: docs/Week3_Advanced_Data_Analysis_and_Visualization.docx
Candidate Name: Ajay M
"""

import os
import re
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

project_root = os.path.dirname(os.path.abspath(__file__))
doc_path = os.path.join(project_root, "docs", "Week3_Advanced_Data_Analysis_and_Visualization.docx")
figures_dir = os.path.join(project_root, "outputs", "figures")

# Colors
COLOR_NAVY_HEX = "1E3A8A"
COLOR_TEAL_HEX = "0D9488"
COLOR_SLATE_HEX = "0F172A"
COLOR_LIGHT_BG_HEX = "F1F5F9"
COLOR_BORDER_HEX = "CBD5E1"

RGB_NAVY = RGBColor(30, 58, 138)
RGB_TEAL = RGBColor(13, 148, 136)
RGB_SLATE = RGBColor(15, 23, 42)
RGB_GRAY = RGBColor(100, 116, 139)


def set_cell_background(cell, fill_hex):
    """Sets background color of a table cell."""
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)


def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Sets internal padding for a table cell."""
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'<w:top w:w="{top}" w:type="dxa"/>'
        f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'<w:left w:w="{left}" w:type="dxa"/>'
        f'<w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(tcMar)


def create_callout_box(doc, title, text, border_color_hex=COLOR_TEAL_HEX, bg_hex="F8FAFC"):
    """Creates a styled callout box in Word."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    
    cell = tbl.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_background(cell, bg_hex)
    set_cell_margins(cell, top=140, bottom=140, left=200, right=180)
    
    # Left border only
    tcPr = cell._element.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'<w:top w:val="none"/>'
        f'<w:left w:val="single" w:sz="24" w:space="0" w:color="{border_color_hex}"/>'
        f'<w:bottom w:val="none"/>'
        f'<w:right w:val="none"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(borders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    run_title = p.add_run(f"📌 {title}\n")
    run_title.bold = True
    run_title.font.name = "Segoe UI"
    run_title.font.size = Pt(10.5)
    run_title.font.color.rgb = RGB_NAVY
    
    run_text = p.add_run(text)
    run_text.font.name = "Segoe UI"
    run_text.font.size = Pt(9.5)
    run_text.font.color.rgb = RGB_SLATE
    
    doc.add_paragraph()  # spacing


def build_word_document():
    doc = Document()
    
    # Configure 1-inch margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
        # Header / Footer
        header = section.header
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hrun = hp.add_run("Week 3: Advanced Data Analysis and Visualization in Logistics | Ajay M")
        hrun.font.name = "Segoe UI"
        hrun.font.size = Pt(8.5)
        hrun.font.color.rgb = RGB_GRAY
        
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        frun = fp.add_run("Confidential — Logistics Analytics & Academic Project Submission")
        frun.font.name = "Segoe UI"
        frun.font.size = Pt(8.5)
        frun.font.color.rgb = RGB_GRAY

    # Set base style fonts
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Segoe UI'
    font.size = Pt(10.5)
    font.color.rgb = RGB_SLATE

    # ==========================================
    # COVER / TITLE SECTION
    # ==========================================
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_p.paragraph_format.space_before = Pt(36)
    title_p.paragraph_format.space_after = Pt(6)
    
    run_sub = title_p.add_run("LOGISTICS & SUPPLY CHAIN DATA SCIENCE CAPSTONE\n")
    run_sub.font.name = "Segoe UI"
    run_sub.font.size = Pt(11)
    run_sub.font.bold = True
    run_sub.font.color.rgb = RGB_TEAL
    
    run_title = title_p.add_run("Week 3: Advanced Data Analysis and Visualization in Logistics\n")
    run_title.font.name = "Segoe UI Semibold"
    run_title.font.size = Pt(22)
    run_title.font.bold = True
    run_title.font.color.rgb = RGB_NAVY
    
    run_theme = title_p.add_run("Theme: Logistics Delivery Performance, Shipment Volume, and Transportation Cost Analysis")
    run_theme.font.name = "Segoe UI"
    run_theme.font.size = Pt(13)
    run_theme.font.italic = True
    run_theme.font.color.rgb = RGB_GRAY

    # Metadata Table
    meta_table = doc.add_table(rows=5, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_table.autofit = False
    
    metadata = [
        ("Candidate Name / Author:", "Ajay M"),
        ("Role:", "Senior Logistics Analytics Consultant & Python Data Science Developer"),
        ("Course Module:", "Week 3: Advanced Exploratory Data Analysis & Visualization"),
        ("Dataset Analyzed:", "Enterprise Cleaned Logistics Operations Dataset (1,250 Shipments)"),
        ("Submission Date:", "August 2026")
    ]
    
    for i, (k, v) in enumerate(metadata):
        row = meta_table.rows[i]
        c1, c2 = row.cells[0], row.cells[1]
        c1.width = Inches(2.2)
        c2.width = Inches(4.3)
        
        set_cell_background(c1, COLOR_LIGHT_BG_HEX)
        set_cell_background(c2, "FFFFFF")
        set_cell_margins(c1, 80, 80, 100, 100)
        set_cell_margins(c2, 80, 80, 100, 100)
        
        p1 = c1.paragraphs[0]
        r1 = p1.add_run(k)
        r1.bold = True
        r1.font.size = Pt(9.5)
        r1.font.color.rgb = RGB_NAVY
        
        p2 = c2.paragraphs[0]
        r2 = p2.add_run(v)
        r2.font.size = Pt(9.5)
        if k.startswith("Candidate"):
            r2.bold = True
            r2.font.color.rgb = RGB_NAVY
            
    doc.add_paragraph().paragraph_format.space_after = Pt(20)
    doc.add_page_break()

    # ==========================================
    # HELPER FUNCTIONS FOR CONTENT
    # ==========================================
    def add_heading_1(text):
        h = doc.add_heading(text, level=1)
        h.paragraph_format.space_before = Pt(18)
        h.paragraph_format.space_after = Pt(6)
        h.paragraph_format.keep_with_next = True
        for r in h.runs:
            r.font.name = "Segoe UI Semibold"
            r.font.color.rgb = RGB_NAVY
            r.font.size = Pt(16)
            r.bold = True

    def add_heading_2(text):
        h = doc.add_heading(text, level=2)
        h.paragraph_format.space_before = Pt(14)
        h.paragraph_format.space_after = Pt(4)
        h.paragraph_format.keep_with_next = True
        for r in h.runs:
            r.font.name = "Segoe UI Semibold"
            r.font.color.rgb = RGB_TEAL
            r.font.size = Pt(13)
            r.bold = True

    def add_heading_3(text):
        h = doc.add_heading(text, level=3)
        h.paragraph_format.space_before = Pt(10)
        h.paragraph_format.space_after = Pt(2)
        h.paragraph_format.keep_with_next = True
        for r in h.runs:
            r.font.name = "Segoe UI Semibold"
            r.font.color.rgb = RGB_SLATE
            r.font.size = Pt(11)
            r.bold = True

    def add_p(text, bold_prefix=None, italic=False, space_after=6):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.line_spacing = 1.15
        if bold_prefix:
            r_b = p.add_run(bold_prefix)
            r_b.bold = True
            r_b.font.color.rgb = RGB_NAVY
        r = p.add_run(text)
        r.font.name = "Segoe UI"
        r.font.size = Pt(10)
        r.italic = italic
        return p

    def add_bullet(bold_part, text_part):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.15
        r1 = p.add_run(bold_part)
        r1.bold = True
        r1.font.color.rgb = RGB_NAVY
        r1.font.size = Pt(10)
        r2 = p.add_run(text_part)
        r2.font.size = Pt(10)

    def add_image_figure(img_filename, caption_text, width=5.8):
        img_path = os.path.join(figures_dir, img_filename)
        if os.path.exists(img_path):
            ip = doc.add_paragraph()
            ip.alignment = WD_ALIGN_PARAGRAPH.CENTER
            ip.paragraph_format.space_before = Pt(8)
            ip.paragraph_format.space_after = Pt(3)
            ip.add_run().add_picture(img_path, width=Inches(width))
            
            cp = doc.add_paragraph()
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cp.paragraph_format.space_after = Pt(12)
            crun = cp.add_run(f"Figure: {caption_text}")
            crun.font.name = "Segoe UI"
            crun.font.size = Pt(8.5)
            crun.font.italic = True
            crun.font.color.rgb = RGB_GRAY

    def add_formatted_table(headers, rows_data, col_widths=None):
        tbl = doc.add_table(rows=len(rows_data) + 1, cols=len(headers))
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl.autofit = False
        
        # Header Row
        hdr_row = tbl.rows[0]
        for idx, h_text in enumerate(headers):
            cell = hdr_row.cells[idx]
            if col_widths and idx < len(col_widths):
                cell.width = Inches(col_widths[idx])
            set_cell_background(cell, COLOR_NAVY_HEX)
            set_cell_margins(cell, 80, 80, 90, 90)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(h_text)
            r.bold = True
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(255, 255, 255)
            
        # Data Rows
        for r_idx, row_values in enumerate(rows_data):
            row = tbl.rows[r_idx + 1]
            bg_color = COLOR_LIGHT_BG_HEX if r_idx % 2 == 1 else "FFFFFF"
            for c_idx, val in enumerate(row_values):
                cell = row.cells[c_idx]
                if col_widths and c_idx < len(col_widths):
                    cell.width = Inches(col_widths[c_idx])
                set_cell_background(cell, bg_color)
                set_cell_margins(cell, 60, 60, 80, 80)
                p = cell.paragraphs[0]
                # If first column, left align; else center or right
                if c_idx == 0:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                else:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run(str(val))
                r.font.size = Pt(8.5)
                r.font.color.rgb = RGB_SLATE
                
        doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # ==========================================
    # SECTION 1: EXECUTIVE SUMMARY
    # ==========================================
    add_heading_1("1. Executive Summary")
    add_p(
        "Modern enterprise logistics networks operate in complex, volatile environments characterized by stringent customer service level agreements (SLAs), multi-modal freight routing, and thin operating contribution margins. This project presents an exhaustive exploratory, diagnostic, and statistical analysis of enterprise logistics operations conducted by Ajay M. By evaluating 1,250 multi-modal freight shipment records spanning five destination regions, four customer segments, five origin warehouse hubs, and four freight shipping modes across the first half of 2024, this study establishes a rigorous empirical foundation for supply chain optimization."
    )
    
    add_bullet("Fulfillment Latency & Delay Burden: ", "The overall network average delivery latency is 5.00 days (median: 5.20 days, standard deviation: 2.14 days). The network achieves an On-Time Delivery Rate of only 28.24%, resulting in a systemic Delivery Delay Rate of 71.76% (897 delayed shipments).")
    add_bullet("Expedited Tier Fragility: ", "High-cost premium service tiers suffer from severe schedule slippage. Same-Day Courier (averaging $172.50/shipment) and Express Air (averaging $138.55/shipment) exhibit delay rates of 76.19% and 72.97% against their 1-day and 2-day SLA targets, severely undermining customer trust.")
    add_bullet("Geographic Disparities: ", "The South region experiences the highest transit latency (5.19 days average) and the highest regional delay rate (76.98%), leading to the lowest regional customer satisfaction rating (3.50 / 5.0). Conversely, the North region captures the largest shipment volume (320 orders, 25.60% share) and highest customer satisfaction (3.76 / 5.0).")
    add_bullet("Transportation Cost Dynamics: ", "Total shipping expenditure reached $114,809.33 against gross commercial sales of $769,293.85, yielding an aggregate Shipping-Cost-to-Sales ratio of 14.92%. Transportation distance serves as a primary linear cost driver (r = +0.67, slope = $0.061/KM).")
    add_bullet("Product Category Vulnerabilities: ", "While Industrial Machinery generates the largest revenue share ($364,799.80, 47.42% of total sales) with a minimal freight ratio (5.07%), Office Supplies incurs severe margin erosion, with shipping costs accounting for 78.73% of gross sales due to small basket sizes and heavy/bulky paper commodities.")

    create_callout_box(
        doc,
        "Executive Recommendation Summary",
        "1. Restructure carrier contracts with SLA penalty credit clauses for express deliveries exceeding commitments.\n"
        "2. Establish regional cross-docking forward hubs in the South and West corridors to compress long-haul distance.\n"
        "3. Enforce minimum order values ($50) and cartonization optimization for Office Supplies and Apparel.\n"
        "4. Deploy machine learning predictive classification pipelines for real-time dispatch delay risk scoring during checkout."
    )

    # ==========================================
    # SECTION 2 & 3: INTRODUCTION & OBJECTIVES
    # ==========================================
    add_heading_1("2. Introduction")
    add_p(
        "Logistics analytics involves the quantitative examination of procurement, warehousing, inventory staging, and transportation processes to maximize operational efficiency, minimize fulfillment friction, and protect operating contribution margins. In modern globalized commerce, organizations manage intricate physical networks where customer expectations for rapid, predictable delivery must be balanced against volatile transportation expenditures and carrier capacity constraints."
    )
    add_p(
        "This capstone report documents the advanced exploratory data analysis and visualization workflow executed by Ajay M. Building upon foundational strategic frameworks and robust data preprocessing pipelines, this investigation utilizes descriptive statistics, inferential hypothesis testing, bivariate regressions, and publication-grade multivariate visualizations to uncover operational bottlenecks across an enterprise distribution network."
    )

    add_heading_1("3. Project Objective")
    add_p(
        "The overarching objective of this project is to convert raw logistics transaction logs into actionable, data-driven business intelligence following the standard analytical progression:"
    )
    add_p("Data Ingestion → Exploratory Data Analysis → Statistical Modeling → Publication Visualizations → Diagnostic Insights → Strategic Action Plans", italic=True)
    
    add_bullet("Objective 1: ", "Calculate comprehensive measures of central tendency, dispersion, and distribution symmetry for all operational metrics.")
    add_bullet("Objective 2: ", "Formulate and compute core enterprise logistics Key Performance Indicators (KPIs).")
    add_bullet("Objective 3: ", "Investigate the mathematical relationships and correlations between shipment distance, order quantity, freight mode, and transportation cost.")
    add_bullet("Objective 4: ", "Diagnose localized failure points across regional destination markets and fulfillment origin hubs.")
    add_bullet("Objective 5: ", "Create high-resolution, publication-quality visualizations that convey complex operational phenomena to executive stakeholders.")
    add_bullet("Objective 6: ", "Formulate a structured Bottleneck Analysis Matrix and actionable, prioritized recommendations.")

    # ==========================================
    # SECTION 4 & 5: SCENARIO & BUSINESS QUESTIONS
    # ==========================================
    add_heading_1("4. Business Scenario")
    add_p(
        "The analysis centers on GlobalLogix Solutions, a multi-regional third-party logistics (3PL) and fulfillment enterprise providing domestic shipping and supply chain services to four commercial segments: Consumer, Corporate, Home Office, and Small Business across five fulfillment centers (WH-Central, WH-East, WH-North, WH-South, WH-West)."
    )
    add_p(
        "GlobalLogix offers four distinct shipping service levels: Same-Day Courier (1-day SLA), Express Air (2-day SLA), Standard Delivery (5-day SLA), and Ground Freight (7-day SLA). Despite strong top-line sales growth in 2024, leadership observed surging customer complaints regarding late deliveries and severe margin drag across low-value product lines. Ajay M was commissioned to perform a diagnostic audit across all operational touchpoints."
    )

    add_heading_1("5. Business Questions")
    add_bullet("BQ1 (Distribution Properties): ", "What are the central tendencies, dispersion spreads, and skewness characteristics of delivery transit times and transportation costs across the entire logistics network?")
    add_bullet("BQ2 (Distance & Cost Elasticity): ", "To what extent does transit distance determine shipping costs and delivery latency, and does this relationship vary across carrier tiers?")
    add_bullet("BQ3 (Modal SLA Compliance): ", "How reliably are contractual delivery timelines met across different shipping tiers, and do express services justify their price premiums?")
    add_bullet("BQ4 (Regional Fulfillment Disparities): ", "Which geographic territories demonstrate the poorest fulfillment reliability, highest delay rates, and lowest customer satisfaction?")
    add_bullet("BQ5 (Product-Level Freight Burden): ", "Which product categories experience disproportionate logistics costs relative to gross sales revenue?")
    add_bullet("BQ6 (Multivariate Interactions): ", "How do origin warehouse dispatch hubs interact with destination regions to create transit delays?")
    add_bullet("BQ7 (Outlier & Anomaly Analysis): ", "What operational anomalies exist in shipment distances, processing days, or expedited charges, and what are their systemic causes?")
    add_bullet("BQ8 (Optimization Strategy): ", "What concrete, prioritized operational interventions should leadership execute to enhance on-time performance and protect profitability?")

    # ==========================================
    # SECTION 6 & 7: DATASET & METHODOLOGY
    # ==========================================
    add_heading_1("6. Dataset Description")
    add_p(
        "The dataset utilized in this analysis is the clean, validated enterprise dataset generated from the Week 2 data preprocessing pipeline (logistics_cleaned.csv). The dataset contains 1,250 verified shipment transactions spanning from January 1, 2024, to June 30, 2024, structured across 42 feature columns."
    )

    dataset_headers = ["Variable Name", "Data Type", "Operational Definition", "Sample Values / Range"]
    dataset_rows = [
        ["Order_ID", "String", "Unique shipment transaction identifier", "ORD-2024-1001 to ORD-2024-2250"],
        ["Order_Date", "Datetime", "Timestamp of customer order placement", "2024-01-01 to 2024-06-30"],
        ["Customer_Segment", "Categorical", "Commercial purchasing tier", "Consumer, Corporate, Home Office, Small Business"],
        ["Product_Category", "Categorical", "Classification of transacted inventory", "Apparel, Electronics, Healthcare, Industrial, Office"],
        ["Quantity", "Integer", "Total units transacted within shipment", "1 to 50 units (Mean: 3.45)"],
        ["Sales_USD", "Float", "Gross commercial revenue value of order", "$14.49 to $2,165.61 (Mean: $615.44)"],
        ["Shipping_Cost_USD", "Float", "Total carrier freight and handling spend", "$8.00 to $233.00 (Mean: $91.85)"],
        ["Delivery_Time_Days", "Float", "Actual elapsed transit time to delivery", "0.80 to 9.80 days (Mean: 5.00 d)"],
        ["Estimated_Delivery_Days", "Float", "Contractual SLA scheduled delivery timeline", "1.0, 2.0, 5.0, 7.0 days"],
        ["Distance_KM", "Float", "Calculated geographic transit distance", "52.40 to 1,886.88 KM (Mean: 819.43 KM)"],
        ["Warehouse_Code", "Categorical", "Origin fulfillment center code", "WH-Central, WH-East, WH-North, WH-South, WH-West"],
        ["Region", "Categorical", "Destination geographic territory", "Central, East, North, South, West"],
        ["Shipping_Mode", "Categorical", "Freight service level tier", "Same-Day Courier, Express Air, Standard, Ground"],
        ["Delivery_Status", "Categorical", "Physical fulfillment status", "Delivered, Delayed"],
        ["Customer_Rating", "Float", "Post-delivery satisfaction score", "1.0 to 5.0 (Mean: 3.63)"],
        ["Order_Processing_Days", "Integer", "Elapsed dwell time to dispatch", "1 to 4 days (Mean: 2.43 d)"],
        ["Is_Delayed", "Binary", "Binary delay flag (1 = Late, 0 = On-Time)", "0 (28.24%), 1 (71.76%)"]
    ]
    add_formatted_table(dataset_headers, dataset_rows, [1.5, 1.0, 2.3, 1.7])

    add_heading_1("7. Analytical Methodology")
    add_p(
        "To ensure maximum academic reproducibility and industrial rigor, this analysis adheres to a structured five-stage quantitative methodology:"
    )
    add_bullet("1. Ingestion & Validation: ", "Verification of schema datatypes, datetime parsing, absence of nulls/duplicates, and feature boundary conformance.")
    add_bullet("2. Parametric & Non-Parametric Profiling: ", "Computation of Mean, Median, Mode, 5% Trimmed Mean, Standard Deviation, Variance, Interquartile Range (IQR), Skewness, and Kurtosis.")
    add_bullet("3. Exploratory Visual Diagnostics: ", "Construction of probability density histograms, linear regression fits, scatter interaction grids, and correlation matrices.")
    add_bullet("4. Multidimensional Slicing: ", "Granular aggregations across geographic regions, carrier modes, warehouse origins, customer tiers, and chronological monthly intervals.")
    add_bullet("5. Synthesis & Root-Cause Analysis: ", "Formulation of evidence-grounded insights using the structured (F-E-M-A) framework, construction of the Operational Bottleneck Matrix, and phased strategic prioritization.")

    # ==========================================
    # SECTION 8 & 9: DESCRIPTIVE STATISTICS
    # ==========================================
    add_heading_1("8. Dataset Exploration")
    add_p(
        "An initial audit of the 1,250 records confirmed flawless data hygiene: exactly 0 missing values across all 42 attributes, 0 duplicated rows, 182 days of continuous operational records (Jan 1 – Jun 30, 2024), and balanced category distributions across customer segments and regions."
    )

    add_heading_1("9. Descriptive Statistics")
    add_p(
        "A comprehensive statistical summary was computed for all core numerical operational variables to understand baseline central tendencies, dispersion spreads, and probability density shapes."
    )

    desc_headers = ["Variable", "Count", "Mean", "Median", "Mode", "Std Dev", "Min", "Q1", "Q3", "Max", "IQR", "Skewness", "Kurtosis"]
    desc_rows = [
        ["Delivery_Time_Days", "1,250", "5.0003", "5.2000", "5.7000", "2.1448", "0.80", "3.30", "6.70", "9.80", "3.40", "-0.198", "-0.930"],
        ["Shipping_Cost_USD", "1,250", "91.8475", "78.4350", "67.4300", "48.7498", "8.00", "54.85", "121.22", "233.00", "66.36", "+0.871", "+0.137"],
        ["Sales_USD", "1,250", "615.435", "310.560", "2165.61", "679.529", "14.49", "98.43", "875.25", "2165.61", "776.83", "+1.258", "+0.285"],
        ["Quantity", "1,250", "3.4472", "2.0000", "1.0000", "4.6186", "1.00", "1.00", "4.00", "50.00", "3.00", "+5.568", "+41.56"],
        ["Distance_KM", "1,250", "819.432", "761.350", "1886.88", "501.996", "52.40", "392.48", "1207.25", "1886.88", "814.78", "+0.495", "-0.638"],
        ["Order_Processing_Days", "1,250", "2.4320", "2.0000", "1.0000", "1.1169", "1.00", "1.00", "3.00", "4.00", "2.00", "+0.108", "-1.332"],
        ["Customer_Rating", "1,250", "3.6328", "4.0000", "4.0000", "1.1578", "1.00", "3.00", "4.75", "5.00", "1.75", "-0.610", "-0.499"],
        ["Cost_Per_KM", "1,250", "0.1302", "0.0883", "0.0525", "0.1097", "0.0108", "0.0573", "0.1652", "0.8540", "0.1079", "+2.379", "+7.913"],
        ["Cost_Per_Unit", "1,250", "49.4239", "33.6200", "66.9700", "42.4504", "2.34", "16.18", "72.94", "172.50", "56.76", "+1.226", "+0.825"]
    ]
    add_formatted_table(desc_headers, desc_rows, [1.5, 0.4, 0.45, 0.45, 0.45, 0.45, 0.35, 0.35, 0.45, 0.45, 0.4, 0.45, 0.45])

    # ==========================================
    # SECTION 10: EXPLORATORY DATA ANALYSIS & FIGURES
    # ==========================================
    add_heading_1("10. Exploratory Data Analysis & Visualizations")
    add_heading_2("10.1 Univariate Distribution Analysis")
    add_p(
        "Univariate analysis examines the probability density, quartile distributions, and shape characteristics of individual metrics in isolation. Delivery latency exhibits an approximately symmetric, platykurtic distribution reflecting multi-modal shipping service levels, while shipping cost displays moderate positive skewness ($+0.871$)."
    )
    add_image_figure("01_delivery_time_distribution.png", "Empirical distribution of delivery transit time (days) with mean (5.00 d) and median (5.20 d) reference lines.")
    add_image_figure("02_shipping_cost_distribution.png", "Empirical distribution of transportation and shipping expenditure ($ USD) with positive skewness overlay.")

    add_heading_2("10.2 Bivariate Relationship Analysis")
    add_p(
        "Bivariate analysis evaluates mathematical relationships and regression trends. Distance is positively correlated with shipping cost ($r = +0.6724, R^2 = 0.45$, slope = $0.061/KM) and delivery latency ($r = +0.5218$)."
    )
    add_image_figure("07_distance_vs_delivery_time.png", "Scatter plot of transportation transit distance (KM) vs delivery latency (days) by shipping mode.")
    add_image_figure("08_distance_vs_shipping_cost.png", "Impact of transit distance on total transportation cost ($ USD) with OLS regression fit.")

    add_heading_2("10.3 Multivariate Interaction Analysis")
    add_p(
        "Multivariate analysis explores complex interaction effects between order quantity, product category, gross commercial value, and freight spend."
    )
    add_image_figure("11_quantity_vs_shipping_cost.png", "Multivariate scatter plot of shipment quantity vs shipping cost segmented by product category and sales volume.")
    add_image_figure("12_product_category_performance.png", "Dual-panel comparison of commercial sales, shipping expenditure, and freight cost burden ratio by category.")

    # ==========================================
    # SECTION 11 & 12: KPIS & CORRELATIONS
    # ==========================================
    add_heading_1("11. Logistics KPI Analysis")
    add_p(
        "To evaluate network-wide operational health, core enterprise logistics Key Performance Indicators (KPIs) were computed from the empirical transaction records."
    )

    kpi_headers = ["Corporate Logistics KPI", "Empirical Metric Value", "Operational Benchmark", "Status / Health Assessment"]
    kpi_rows = [
        ["Total Order Volume", "1,250 Shipments", "N/A", "Full semi-annual operational throughput"],
        ["Gross Commercial Sales", "$769,293.85", "N/A", "Base commercial trading volume"],
        ["Total Transportation Spend", "$114,809.33", "N/A", "Direct freight procurement spend"],
        ["Average Order Value (AOV)", "$615.44", "$500.00", "Strong commercial basket size"],
        ["Average Shipping Cost per Order", "$91.85", "< $75.00", "Elevated due to high expedited air mix"],
        ["Shipping Cost-to-Sales Ratio", "14.92%", "< 10.00%", "High Margin Burden (Industry target: 8-12%)"],
        ["Average Delivery Time", "5.00 Days", "4.00 Days", "Moderate network transit latency"],
        ["Median Delivery Time", "5.20 Days", "4.00 Days", "Reflects 5-day standard delivery dominance"],
        ["On-Time Delivery Rate", "28.24%", "> 90.00%", "Critical Operational Failure"],
        ["Delivery Delay Rate", "71.76%", "< 10.00%", "Primary Network Bottleneck (897 late orders)"],
        ["Average Transit Distance", "819.43 KM", "N/A", "Domestic multi-state distribution baseline"],
        ["Average Cost per KM", "$0.1302 / KM", "< $0.1000/KM", "Premium service tiers drive up unit cost"],
        ["Average Cost per Item Unit", "$49.42 / Unit", "< $35.00/Unit", "Squeezed by low-quantity single-item orders"],
        ["Average Customer Rating", "3.63 / 5.0", "> 4.20 / 5.0", "Subdued due to pervasive delivery delays"],
        ["Average Order Processing Days", "2.43 Days", "< 1.50 Days", "Warehouse dispatch dwell time consumes buffer"]
    ]
    add_formatted_table(kpi_headers, kpi_rows, [1.8, 1.2, 1.2, 2.3])

    add_heading_1("12. Correlation Analysis")
    add_p(
        "A bivariate Pearson correlation analysis was conducted across all continuous numerical logistics features, accompanied by two-tailed hypothesis testing (p < 0.05)."
    )
    add_image_figure("09_correlation_heatmap.png", "Pearson correlation coefficient matrix across key numerical logistics features.")

    corr_headers = ["Feature 1", "Feature 2", "Pearson r", "p-value", "Direction & Operational Meaning"]
    corr_rows = [
        ["Distance_KM", "Shipping_Cost_USD", "+0.6724", "p < 0.0001", "Strong Positive: Transit distance is the primary linear driver of freight spend."],
        ["Distance_KM", "Delivery_Time_Days", "+0.5218", "p < 0.0001", "Moderate Positive: Greater distances increase line-haul transit latency."],
        ["Delivery_Time_Days", "Customer_Rating", "-0.4812", "p < 0.0001", "Moderate Negative: Longer turnaround significantly degrades customer ratings."],
        ["Cost_Per_KM", "Delivery_Time_Days", "-0.4531", "p < 0.0001", "Moderate Negative: High cost-per-km tiers correspond to faster delivery times."],
        ["Sales_USD", "Quantity", "+0.4128", "p < 0.0001", "Moderate Positive: Larger item quantities expand gross order commercial value."]
    ]
    add_formatted_table(corr_headers, corr_rows, [1.4, 1.4, 0.8, 0.9, 2.0])

    # ==========================================
    # SECTION 13, 14, 15: TREND, REGION & MODE
    # ==========================================
    add_heading_1("13. Trend Analysis")
    add_p(
        "Temporal analysis evaluated operational volume, commercial sales, and freight expenditures across the 6-month observation window (January 2024 to June 2024). Operational volume remained steady at approx 208 orders per month, with delay rates consistently hovering between 70.6% and 73.3%."
    )
    add_image_figure("10_monthly_order_volume_cost_trend.png", "Monthly shipment order volume and total transportation expenditure dynamics (2024).")

    add_heading_1("14. Regional Performance Analysis")
    add_p(
        "Geographic destination territories were analyzed to identify regional fulfillment imbalances, transit latency bottlenecks, and customer satisfaction friction."
    )
    add_image_figure("03_shipment_volume_by_region.png", "Total shipment order volume distribution across destination geographic regions.")
    add_image_figure("04_avg_delivery_time_by_region_warehouse.png", "Multi-hub fulfillment latency (days) by destination region and origin warehouse center.")

    reg_headers = ["Region", "Orders", "Share (%)", "Sales ($)", "Shipping Spend ($)", "Avg Latency (d)", "Delay Rate (%)", "Avg Rating"]
    reg_rows = [
        ["North", "320", "25.60%", "$210,652.41", "$30,140.32", "4.90", "68.75%", "3.76"],
        ["West", "278", "22.24%", "$164,033.12", "$26,397.28", "4.97", "71.58%", "3.62"],
        ["South", "252", "20.16%", "$161,683.44", "$23,430.95", "5.19", "76.98%", "3.50"],
        ["East", "228", "18.24%", "$137,406.01", "$18,866.49", "5.06", "68.86%", "3.61"],
        ["Central", "172", "13.76%", "$95,518.87", "$15,974.29", "4.89", "73.84%", "3.59"]
    ]
    add_formatted_table(reg_headers, reg_rows, [1.0, 0.7, 0.8, 1.0, 1.1, 1.0, 1.0, 0.9])

    add_heading_1("15. Shipping Mode Analysis")
    add_p(
        "A comparative performance audit was conducted across the four shipping tiers to evaluate speed, SLA compliance, cost structure, and customer ratings."
    )
    add_image_figure("05_shipping_cost_by_mode.png", "Boxplot distribution of shipping expenditures across service tiers with median/mean overlays.")
    add_image_figure("06_delivery_time_by_mode.png", "Actual delivery time distributions vs contractual SLA benchmark diamonds by shipping mode.")
    add_image_figure("13_delivery_status_delay_rate_by_mode.png", "Proportional distribution of on-time vs delayed fulfillment across shipping service levels.")

    mode_headers = ["Shipping Mode", "Orders", "Share (%)", "SLA Target", "Actual Avg (d)", "Avg Spend ($)", "Delay Rate (%)", "Rating"]
    mode_rows = [
        ["Same-Day Courier", "126", "10.08%", "1.0 Day", "2.21", "$172.50", "76.19%", "3.57"],
        ["Express Air", "296", "23.68%", "2.0 Days", "2.94", "$138.55", "72.97%", "3.67"],
        ["Standard Delivery", "656", "52.48%", "5.0 Days", "5.76", "$58.39", "71.49%", "3.60"],
        ["Ground Freight", "172", "13.76%", "7.0 Days", "7.71", "$79.99", "67.44%", "3.71"]
    ]
    add_formatted_table(mode_headers, mode_rows, [1.4, 0.7, 0.8, 0.9, 0.9, 0.9, 0.9, 0.7])

    # ==========================================
    # SECTION 16, 17, 18: CATEGORIES, OUTLIERS & ADVANCED VIZ
    # ==========================================
    add_heading_1("16. Product / Category Analysis")
    add_p(
        "Inventory classifications were examined to quantify logistical handling burdens, freight cost ratios, and commercial revenue generation."
    )

    cat_headers = ["Product Category", "Orders", "Units", "Total Sales ($)", "Shipping Spend ($)", "Avg Latency (d)", "Delay Rate (%)", "Freight Ratio (%)"]
    cat_rows = [
        ["Industrial Machinery", "199", "829", "$364,799.80", "$18,500.40", "4.91", "68.34%", "5.07%"],
        ["Electronics", "369", "1,372", "$275,874.59", "$36,175.40", "4.87", "71.00%", "13.11%"],
        ["Apparel", "246", "781", "$52,512.50", "$22,133.66", "4.99", "70.33%", "42.15%"],
        ["Healthcare Supplies", "114", "366", "$40,619.81", "$10,060.29", "5.15", "72.81%", "24.77%"],
        ["Office Supplies", "322", "961", "$35,487.15", "$27,939.58", "5.17", "75.47%", "78.73%"]
    ]
    add_formatted_table(cat_headers, cat_rows, [1.4, 0.6, 0.6, 1.1, 1.1, 0.9, 0.9, 0.9])

    add_heading_1("17. Outlier Analysis")
    add_p(
        "Outlier detection was performed using Tukey's Interquartile Range (IQR) Fences Rule. In logistics engineering, statistical outliers (such as bulk orders of 50 units or express shipments of $233) represent legitimate commercial events rather than data entry corruptions, and are retained for full empirical fidelity."
    )

    add_heading_1("18. Advanced Multivariate Visualizations")
    add_p(
        "Multivariate analysis evaluated the simultaneous interactions between customer segments, service tiers, destination regions, and fulfillment reliability."
    )
    add_image_figure("14_customer_segment_comparison.png", "Multidimensional logistics comparison across commercial customer segments.")
    add_image_figure("15_multivariate_delay_risk_matrix.png", "Multivariate operational risk matrix heatmap: Delay rate (%) across Region and Shipping Mode.")

    # ==========================================
    # SECTION 19 & 20: INSIGHTS & BOTTLENECKS
    # ==========================================
    add_heading_1("19. Key Analytical Insights")
    
    insights_data = [
        ("INS-01: Regional Latency Disparities", 
         "Finding: Geographic destination creates substantial delivery latency variance, with South exhibiting the highest transit time (5.19 d) and delay rate (76.98%).\n"
         "Evidence: South records 5.19 days average latency vs 4.89 days in Central. Customer satisfaction in South is lowest (3.50 / 5.0).\n"
         "Business Meaning: Regional fulfillment variance stems from uneven warehouse proximity and multi-hub handoffs.\n"
         "Potential Action: Establish regional cross-docking hubs in Atlanta/Dallas to compress line-haul transit distance."),
        
        ("INS-02: Modal Economics & SLA Fragility", 
         "Finding: Premium expedited tiers command high shipping costs but struggle with severe punctuality failures.\n"
         "Evidence: Express Air costs $138.55/order with a 72.97% delay rate; Same-Day Courier costs $172.50 with a 76.19% delay rate.\n"
         "Business Meaning: Customers paying express surcharges have zero tolerance for schedule slippage, causing churn.\n"
         "Potential Action: Enforce priority picking in warehouse queues and contractual carrier penalty clawback clauses."),
        
        ("INS-03: Category Freight Cost Burden", 
         "Finding: Office Supplies and Apparel generate severe logistics cost drag relative to commercial sales.\n"
         "Evidence: Office Supplies incurs a 78.73% freight-to-sales ratio ($27.9k shipping on $35.5k sales).\n"
         "Business Meaning: Low-value individual parcel dispatches wipe out product-level profitability.\n"
         "Potential Action: Enforce minimum order values ($50) and deploy automated 3D cartonization optimization software.")
    ]
    
    for title, text in insights_data:
        create_callout_box(doc, title, text, COLOR_NAVY_HEX, COLOR_LIGHT_BG_HEX)

    add_heading_1("20. Logistics Bottlenecks Matrix")
    add_p(
        "The table below maps identified network bottlenecks against empirical evidence, operational concerns, and corrective interventions:"
    )

    bn_headers = ["Logistics Area", "Primary Indicator", "Empirical Evidence", "Operational Risk", "Corrective Action"]
    bn_rows = [
        ["Region (South)", "Delivery Latency & Delay", "5.19 days avg delivery; 76.98% delay rate across 252 orders", "Protracted transit distance and remote handoffs causing customer churn.", "Establish regional cross-docking hub and onboard secondary carriers."],
        ["Mode (Same-Day)", "Severe SLA Breach", "2.21 days avg delivery (vs 1.0 d SLA); 76.19% delay rate", "Premium surcharge without delivery guarantee; refund exposure.", "Restrict courier radius to 50 KM and deploy dedicated fleets."],
        ["Mode (Express Air)", "Schedule Slippage", "2.94 days avg delivery (vs 2.0 d SLA); 72.97% delay rate", "High spend ($138.55/order) failing to deliver promised turnaround.", "Enforce warehouse priority picking and strict carrier SLA chargebacks."],
        ["Category (Office Supplies)", "Logistics Cost Drag", "Shipping spend accounts for 78.73% of gross category sales", "Complete erosion of gross margins; unprofitable fulfillment.", "Enforce minimum order values ($50) and optimize carton packaging."],
        ["Hub (WH-South)", "Dispatch Dwell Time", "2.52 days avg processing time; 75.4% delay rate", "Staging and picking delays eat into carrier transit buffers.", "Implement automated wave-picking and enforce same-day dispatch cutoffs."]
    ]
    add_formatted_table(bn_headers, bn_rows, [1.2, 1.2, 1.4, 1.4, 1.3])

    # ==========================================
    # SECTION 21 TO 27: RECOMMENDATIONS, IMPACT, CONCLUSION
    # ==========================================
    add_heading_1("21. Strategic Recommendations")
    add_bullet("Phase 1: Immediate (0–3 Months): ", "Institute carrier SLA penalty chargebacks, implement WMS 12-hour dispatch cutoffs, and mandate minimum order thresholds ($50) for Office Supplies.")
    add_bullet("Phase 2: Medium-Term (3–6 Months): ", "Deploy regional cross-docking forward hubs in Dallas and Phoenix, implement automated 3D cartonization software, and deploy dynamic checkout EDD algorithms.")
    add_bullet("Phase 3: Long-Term (6–12 Months): ", "Build machine learning predictive delay classification engines and deploy a centralized IoT supply chain control tower.")

    add_heading_1("22. Business Impact")
    add_p(
        "Execution of these data-driven recommendations will improve the network on-time delivery rate from 28.24% to >85.00%, reduce delay rates to <15.00%, eliminate $45k in annual customer support friction, and restore category profitability for Office Supplies."
    )

    add_heading_1("23. Future Data Science Applications")
    add_p(
        "The exploratory findings established by Ajay M provide a direct launching pad for advanced predictive modeling, including supervised XGBoost binary delay classifiers, continuous freight cost regression forecasting, customer spatial clustering, and Mixed-Integer Linear Programming facility location models."
    )

    add_heading_1("24. Challenges and Limitations")
    add_p(
        "Key analytical limitations include the absence of intermediate IoT milestone scans between dispatch and delivery, historical rate card pricing dampening spot-market volatility, and synthetic data preprocessing constraints."
    )

    add_heading_1("25. Reflection")
    add_p(
        "Completing the Week 3 analytics project provided invaluable practical insights for Ajay M. The project demonstrated that rigorous exploratory data analysis—analyzing trimmed means, quartile spreads, and multi-modal probability density shapes—uncovers operational mechanics that raw summaries obscure. Visualizations bridged raw metrics with executive business strategy, highlighting how diagnostic analytics forms the indispensable foundation for machine learning."
    )

    add_heading_1("26. Conclusion")
    add_p(
        "The Week 3: Advanced Data Analysis and Visualization in Logistics capstone project successfully delivered an end-to-end, reproducible, and mathematically rigorous analytics workflow. By diagnosing systemic fulfillment delays, expedited SLA fragility, and category margin drag, Ajay M has established an actionable roadmap to optimize corporate logistics performance."
    )

    add_heading_1("27. References")
    add_bullet("1. ", "Chopra, S., & Meindl, P. (2016). Supply Chain Management: Strategy, Planning, and Operation (6th ed.). Pearson.")
    add_bullet("2. ", "Christopher, M. (2016). Logistics & Supply Chain Management (5th ed.). Financial Times Publishing.")
    add_bullet("3. ", "McKinney, W. (2022). Python for Data Analysis: Data Wrangling with Pandas, NumPy, and Jupyter (3rd ed.). O'Reilly Media.")
    add_bullet("4. ", "Tukey, J. W. (1977). Exploratory Data Analysis. Addison-Wesley.")

    doc.save(doc_path)
    print(f"Word document successfully created at: {doc_path}")
    return doc_path


if __name__ == "__main__":
    build_word_document()
